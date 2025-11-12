from flask import Flask, render_template, request, jsonify, send_file
from flask_cors import CORS
import os
import shutil
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import io
from datetime import datetime
import os
import pandas as pd
from dotenv import load_dotenv

from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.document_loaders import DataFrameLoader
from langchain_community.vectorstores import FAISS
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate

app = Flask(__name__)
CORS(app)

# Load environment variables
load_dotenv()
api_key = os.getenv("GOOGLE_API_KEY")

# Initialize AI components (same as main.py)
llm = ChatGoogleGenerativeAI(
    model="gemini-2.5-flash",
    temperature=0.5,  # Balanced for natural, flowing language while staying factual
    google_api_key=api_key,
    max_output_tokens=3000  # Allow longer, more detailed reports
)

# Global variables for data management
excel_path = "SECRETARY FORM(1-6).xlsx"
df = None
vectorstore = None
retriever = None
report_chain = None
normal_chain = None

def load_excel_data():
    """Load Excel data and create vector store"""
    global df, vectorstore, retriever, report_chain, normal_chain
    
    if not os.path.exists(excel_path):
        print(f"Warning: {excel_path} not found. Please upload an Excel file.")
        return False
    
    try:
        df = pd.read_excel(excel_path)
        print(f"Loaded Excel file with {len(df)} rows")
        
        # Create documents and vector store
        df['document_content'] = df.apply(prepare_document_content, axis=1)
        loader = DataFrameLoader(df, page_content_column='document_content')
        documents = loader.load()

        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        vectorstore = FAISS.from_documents(documents, embeddings)
        retriever = vectorstore.as_retriever()

        # Create QA chains
        report_chain = RetrievalQA.from_chain_type(
            llm=llm,
            retriever=retriever,
            chain_type="stuff",
            chain_type_kwargs={"prompt": report_prompt}
        )

        normal_chain = RetrievalQA.from_chain_type(
            llm=llm,
            retriever=retriever,
            chain_type="stuff",
            chain_type_kwargs={"prompt": normal_prompt}
        )
        
        return True
    except Exception as e:
        print(f"Error loading Excel data: {e}")
        return False

# Load initial data
load_excel_data()

def prepare_document_content(row):
    """Convert DataFrame row to meaningful text content"""
    content_parts = []
    
    # Add basic info
    if pd.notna(row.get('Name')):
        content_parts.append(f"Name: {row['Name']}")
    if pd.notna(row.get('Category')):
        content_parts.append(f"Category: {row['Category']}")
    
    # Add financial information
    if pd.notna(row.get('Total Collected (Amount)')):
        content_parts.append(f"Total Collected: {row['Total Collected (Amount)']}")
    if pd.notna(row.get('Total Spent (Amount)')):
        content_parts.append(f"Total Spent: {row['Total Spent (Amount)']}")
    if pd.notna(row.get('Total Remaining (Amount)')):
        content_parts.append(f"Total Remaining: {row['Total Remaining (Amount)']}")
    
    # Add event information
    if pd.notna(row.get('Event Name')):
        content_parts.append(f"Event: {row['Event Name']}")
    if pd.notna(row.get('Location')):
        content_parts.append(f"Location: {row['Location']}")
    if pd.notna(row.get('Attendance')):
        content_parts.append(f"Attendance: {row['Attendance']}")
    if pd.notna(row.get('Key Outcomes')):
        content_parts.append(f"Key Outcomes: {row['Key Outcomes']}")
    
    # Add meeting information
    if pd.notna(row.get('Agenda')):
        content_parts.append(f"Agenda: {row['Agenda']}")
    if pd.notna(row.get('Decisions Made')):
        content_parts.append(f"Decisions Made: {row['Decisions Made']}")
    
    # Add issues and comments
    if pd.notna(row.get('Issue Title')):
        content_parts.append(f"Issue: {row['Issue Title']}")
    if pd.notna(row.get('Description')):
        content_parts.append(f"Description: {row['Description']}")
    if pd.notna(row.get('Comments')):
        content_parts.append(f"Comments: {row['Comments']}")
    
    return " | ".join(content_parts) if content_parts else f"Record ID: {row.get('ID', 'Unknown')}"

# Prompt templates
report_template = """
You are the Welfare Committee Secretary writing a warm, human, and inclusive report.

CRITICAL INSTRUCTIONS:
1. Use ONLY the data provided in the context below - DO NOT use general knowledge or make assumptions
2. Write in first person as the secretary, using a warm and inclusive tone
3. Write in flowing paragraphs and essay form - NO bullet points or lists
4. Make it read like a story that connects all the welfare activities together
5. Include ALL specific details from the data: names, amounts, dates, locations, attendance, outcomes

STRUCTURE YOUR REPORT AS FLOWING PARAGRAPHS:

Opening: Start with a warm introduction about the welfare committee's mission and period covered.

Financial Overview: Write a detailed narrative paragraph about the finances. Mention total collections, how the money was spent, what remains, and what this means for the welfare of members. Make it conversational but include all the numbers.

Events and Activities: Write in essay form about each event. Describe what happened, who attended, where it was held, and what outcomes were achieved. Connect the events to show how they contribute to member welfare. Make it feel inclusive and celebratory.

Meetings and Decisions: Write flowing paragraphs about the meetings held. Discuss the agendas naturally, weave in the decisions made, and explain how these decisions impact the welfare of members. Make it read like you're telling a colleague about important discussions.

Challenges and Progress: If there are issues or concerns in the data, discuss them in a supportive, solution-oriented way. Show empathy and commitment to addressing them.

Closing: End with a warm, forward-looking statement about the committee's continued commitment to member welfare.

TONE: Warm, professional, inclusive, human, and caring. Avoid corporate jargon. Write as if you're sharing good news with friends while maintaining professionalism.

User request: {question}

Welfare Committee Data:
{context}

Write your report in flowing essay form with connected paragraphs. NO bullet points. Make it human and inclusive.
"""

normal_template = """
You are the Welfare Committee Secretary answering a specific question in a warm, human way.

CRITICAL INSTRUCTIONS:
1. Use ONLY the data provided in the context below
2. Write in a conversational, friendly tone while being professional
3. Include specific details from the data: names, amounts, dates, locations
4. Write in flowing sentences, not bullet points
5. If the data doesn't contain the answer, say "I don't have that information in our current records"
6. Make your response feel personal and caring, as befits a welfare committee

User request: {question}

Welfare Committee Data:
{context}

Answer warmly and conversationally using only the data above:
"""

report_prompt = PromptTemplate(template=report_template, input_variables=["context", "question"])
normal_prompt = PromptTemplate(template=normal_template, input_variables=["context", "question"])

# QA chains will be created in load_excel_data()

def generate_response(query):
    global report_chain, normal_chain
    
    if report_chain is None or normal_chain is None:
        return "Please upload an Excel file first to enable AI responses."
    
    if "report" in query.lower():
        chain_to_use = report_chain
    else:
        chain_to_use = normal_chain

    try:
        response = chain_to_use.run(query)
        return response
    except Exception as e:
        return f"An error occurred: {str(e)}"

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/admin')
def admin():
    return render_template('admin.html')

@app.route('/api/upload-excel', methods=['POST'])
def upload_excel():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not file.filename.lower().endswith(('.xlsx', '.xls')):
        return jsonify({'error': 'Please upload an Excel file (.xlsx or .xls)'}), 400
    
    try:
        # Backup current file if it exists
        if os.path.exists(excel_path):
            backup_path = f"{excel_path}.backup"
            shutil.copy2(excel_path, backup_path)
            print(f"Backed up current file to {backup_path}")
        
        # Save new file
        file.save(excel_path)
        print(f"Saved new Excel file: {excel_path}")
        
        # Reload data
        success = load_excel_data()
        if success:
            return jsonify({
                'message': 'Excel file uploaded and processed successfully!',
                'rows': len(df) if df is not None else 0
            })
        else:
            return jsonify({'error': 'Failed to process the Excel file'}), 500
            
    except Exception as e:
        return jsonify({'error': f'Upload failed: {str(e)}'}), 500

@app.route('/api/data-status')
def data_status():
    global df
    if df is not None:
        return jsonify({
            'loaded': True,
            'rows': len(df),
            'columns': list(df.columns),
            'last_modified': os.path.getmtime(excel_path) if os.path.exists(excel_path) else None
        })
    else:
        return jsonify({
            'loaded': False,
            'message': 'No Excel data loaded'
        })

def clean_content_for_document(content):
    """Clean content to remove AI-specific language"""
    clean_content = content.replace("I am", "The committee").replace("I have", "The committee has")
    clean_content = clean_content.replace("my role", "the secretary's role")
    clean_content = clean_content.replace("I can", "The committee can")
    clean_content = clean_content.replace("I will", "The committee will")
    return clean_content

def create_report_pdf(content):
    """Create a clean, professional PDF report without metadata"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, 
                          rightMargin=72, leftMargin=72,
                          topMargin=72, bottomMargin=72)
    
    # Get styles
    styles = getSampleStyleSheet()
    
    # Custom styles for professional report
    title_style = ParagraphStyle(
        'ReportTitle',
        parent=styles['Heading1'],
        fontSize=20,
        spaceAfter=30,
        alignment=TA_CENTER,
        textColor='#1a365d',
        fontName='Helvetica-Bold'
    )
    
    body_style = ParagraphStyle(
        'ReportBody',
        parent=styles['Normal'],
        fontSize=12,
        spaceAfter=14,
        alignment=TA_JUSTIFY,
        leftIndent=0,
        rightIndent=0,
        leading=16
    )
    
    # Build clean PDF content
    story = []
    
    # Clean title
    story.append(Paragraph("Welfare Committee Report", title_style))
    story.append(Spacer(1, 30))
    
    # Process content
    clean_content = clean_content_for_document(content)
    
    # Split content into paragraphs and format professionally
    paragraphs = clean_content.split('\n\n')
    for para in paragraphs:
        if para.strip():
            # Clean up the paragraph
            clean_para = para.strip().replace('\n', ' ')
            story.append(Paragraph(clean_para, body_style))
            story.append(Spacer(1, 8))
    
    # Professional footer with date
    story.append(Spacer(1, 40))
    current_date = datetime.now().strftime("%B %d, %Y")
    story.append(Paragraph(f"Report Date: {current_date}", styles['Normal']))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer

def create_report_docx(content):
    """Create a clean, professional Word document report"""
    buffer = io.BytesIO()
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add title
    title = doc.add_heading('Welfare Committee Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add some space
    doc.add_paragraph()
    
    # Process content
    clean_content = clean_content_for_document(content)
    
    # Split content into paragraphs and add to document
    paragraphs = clean_content.split('\n\n')
    for para in paragraphs:
        if para.strip():
            # Clean up the paragraph
            clean_para = para.strip().replace('\n', ' ')
            p = doc.add_paragraph(clean_para)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Add footer with date
    doc.add_paragraph()
    doc.add_paragraph()
    current_date = datetime.now().strftime("%B %d, %Y")
    footer_para = doc.add_paragraph(f"Report Date: {current_date}")
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Save to buffer
    doc.save(buffer)
    buffer.seek(0)
    return buffer

@app.route('/api/chat', methods=['POST'])
def chat():
    data = request.get_json()
    query = data.get('message', '')
    
    if not query:
        return jsonify({'error': 'No message provided'}), 400
    
    response = generate_response(query)
    
    # Check if this is a report that should have PDF capability
    is_report = (
        "report" in query.lower() or 
        len(response) > 150 or  # Long responses are likely reports
        any(keyword in response.lower() for keyword in 
            ['committee', 'activities', 'finances', 'events', 'meetings', 'summary', 'welfare', 'collected', 'spent'])
    )
    
    print(f"Query: {query}")
    print(f"Response length: {len(response)}")
    print(f"Is report: {is_report}")
    
    return jsonify({
        'response': response,
        'is_report': is_report
    })

@app.route('/api/download-pdf', methods=['POST'])
def download_pdf():
    data = request.get_json()
    response = data.get('response', '')
    
    if not response:
        return jsonify({'error': 'Missing response content'}), 400
    
    try:
        # Create clean PDF report
        pdf_buffer = create_report_pdf(response)
        
        # Generate filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"welfare_committee_report_{timestamp}.pdf"
        
        return send_file(
            pdf_buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/pdf'
        )
    except Exception as e:
        return jsonify({'error': f'PDF generation failed: {str(e)}'}), 500

@app.route('/api/download-docx', methods=['POST'])
def download_docx():
    data = request.get_json()
    response = data.get('response', '')
    
    if not response:
        return jsonify({'error': 'Missing response content'}), 400
    
    try:
        # Create clean Word document report
        docx_buffer = create_report_docx(response)
        
        # Generate filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"welfare_committee_report_{timestamp}.docx"
        
        return send_file(
            docx_buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'error': f'Word document generation failed: {str(e)}'}), 500

@app.route('/test')
def test():
    return jsonify({
        'message': 'Test endpoint working',
        'is_report': True
    })

def find_free_port():
    """Find and return a free port"""
    import socket
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        s.bind(('', 0))
        s.listen(1)
        port = s.getsockname()[1]
    return port

if __name__ == '__main__':
    import os
    # Try to use environment PORT first, otherwise find a free port
    if 'PORT' in os.environ:
        port = int(os.environ.get('PORT'))
    else:
        port = find_free_port()
    
    print(f"Starting Flask app on port {port}")
    print(f"Access your app at: http://localhost:{port}")
    app.run(host='0.0.0.0', port=port, debug=False)