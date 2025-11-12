import streamlit as st
import pandas as pd
import os
from datetime import datetime
import io
from dotenv import load_dotenv
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.document_loaders import DataFrameLoader
from langchain_community.vectorstores import FAISS
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Set page config
st.set_page_config(
    page_title="Welfare Secretary AI",
    page_icon="🏛️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Load environment variables
load_dotenv()
api_key = os.getenv("GOOGLE_API_KEY")

# Custom CSS - Dark Mode Theme
st.markdown("""
<style>
    /* Main app background */
    .stApp {
        background-color: #0a0a0a;
    }
    
    /* Sidebar styling */
    [data-testid="stSidebar"] {
        background-color: #1a1a1a;
        border-right: 1px solid #333;
    }
    
    /* Main header */
    .main-header {
        text-align: center;
        padding: 2rem 1rem;
        background: linear-gradient(135deg, #2d2d2d, #1a1a1a);
        color: #ffffff;
        border-radius: 15px;
        margin-bottom: 2rem;
        border: 1px solid #404040;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.3);
    }
    
    .main-header h1 {
        font-size: 2.5rem;
        font-weight: 700;
        margin-bottom: 0.5rem;
        background: linear-gradient(135deg, #ffffff, #b0b0b0);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
    }
    
    .main-header p {
        color: #a0a0a0;
        font-size: 1.1rem;
    }
    
    /* Chat message containers */
    .chat-message {
        padding: 1.25rem;
        margin: 1rem 0;
        border-radius: 12px;
        border-left: 4px solid;
        animation: fadeIn 0.3s ease-in;
    }
    
    @keyframes fadeIn {
        from { opacity: 0; transform: translateY(10px); }
        to { opacity: 1; transform: translateY(0); }
    }
    
    .user-message {
        background: linear-gradient(135deg, #2a2a2a, #1f1f1f);
        border-left-color: #808080;
        color: #ffffff;
        border: 1px solid #404040;
    }
    
    .ai-message {
        background: linear-gradient(135deg, #1f1f1f, #151515);
        border-left-color: #ffffff;
        color: #e0e0e0;
        border: 1px solid #333333;
    }
    
    .message-label {
        font-weight: 600;
        font-size: 0.9rem;
        margin-bottom: 0.5rem;
        text-transform: uppercase;
        letter-spacing: 1px;
    }
    
    .user-label {
        color: #a0a0a0;
    }
    
    .ai-label {
        color: #ffffff;
    }
    
    .message-content {
        line-height: 1.6;
        font-size: 1rem;
    }
    
    /* Buttons */
    .stButton > button {
        background: linear-gradient(135deg, #2d2d2d, #1a1a1a);
        color: #ffffff;
        border: 1px solid #505050;
        border-radius: 8px;
        padding: 0.5rem 1.5rem;
        font-weight: 500;
        transition: all 0.3s ease;
    }
    
    .stButton > button:hover {
        background: linear-gradient(135deg, #3d3d3d, #2a2a2a);
        border-color: #707070;
        transform: translateY(-2px);
        box-shadow: 0 4px 8px rgba(255, 255, 255, 0.1);
    }
    
    /* Download buttons */
    .stDownloadButton > button {
        background: linear-gradient(135deg, #404040, #2d2d2d);
        color: #ffffff;
        border: 1px solid #606060;
        border-radius: 8px;
        font-size: 0.9rem;
    }
    
    .stDownloadButton > button:hover {
        background: linear-gradient(135deg, #505050, #3d3d3d);
        border-color: #808080;
    }
    
    /* Input field */
    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea {
        background-color: #1a1a1a;
        color: #ffffff;
        border: 1px solid #404040;
        border-radius: 8px;
    }
    
    .stTextInput > div > div > input:focus,
    .stTextArea > div > div > textarea:focus {
        border-color: #808080;
        box-shadow: 0 0 0 1px #808080;
    }
    
    /* Chat input */
    .stChatInput > div {
        background-color: #1a1a1a;
        border: 1px solid #404040;
        border-radius: 12px;
    }
    
    .stChatInput input {
        background-color: #1a1a1a;
        color: #ffffff;
    }
    
    /* File uploader */
    .stFileUploader {
        background-color: #1a1a1a;
        border: 1px dashed #505050;
        border-radius: 8px;
        padding: 1rem;
    }
    
    /* Success/Warning/Error messages */
    .stSuccess {
        background-color: #1a2a1a;
        color: #90ee90;
        border: 1px solid #2d4d2d;
    }
    
    .stWarning {
        background-color: #2a2a1a;
        color: #ffa500;
        border: 1px solid #4d4d2d;
    }
    
    .stError {
        background-color: #2a1a1a;
        color: #ff6b6b;
        border: 1px solid #4d2d2d;
    }
    
    /* Spinner */
    .stSpinner > div {
        border-color: #ffffff transparent transparent transparent;
    }
    
    /* Markdown text */
    .stMarkdown {
        color: #e0e0e0;
    }
    
    /* Headers */
    h1, h2, h3, h4, h5, h6 {
        color: #ffffff !important;
    }
    
    /* Divider */
    hr {
        border-color: #333333;
    }
    
    /* Quick action section */
    .quick-actions {
        background: linear-gradient(135deg, #1f1f1f, #151515);
        padding: 1.5rem;
        border-radius: 12px;
        border: 1px solid #333333;
        margin-top: 2rem;
    }
</style>
""", unsafe_allow_html=True)

# Initialize AI components
@st.cache_resource
def initialize_ai():
    if not api_key:
        return None, None, None
    
    llm = ChatGoogleGenerativeAI(
        model="gemini-2.5-pro",
        temperature=0.5,  # Balanced for natural, flowing language while staying factual
        google_api_key=api_key,
        max_output_tokens=3000  # Allow longer, more detailed reports
    )
    
    # Load Excel data
    excel_path = "SECRETARY FORM(1-6).xlsx"
    if not os.path.exists(excel_path):
        return None, None, None
    
    try:
        with st.spinner("Loading data and initializing AI models..."):
            df = pd.read_excel(excel_path)
        
        # Prepare document content
        def prepare_document_content(row):
            content_parts = []
            
            if pd.notna(row.get('Name')):
                content_parts.append(f"Name: {row['Name']}")
            if pd.notna(row.get('Category')):
                content_parts.append(f"Category: {row['Category']}")
            if pd.notna(row.get('Total Collected (Amount)')):
                content_parts.append(f"Total Collected: {row['Total Collected (Amount)']}")
            if pd.notna(row.get('Total Spent (Amount)')):
                content_parts.append(f"Total Spent: {row['Total Spent (Amount)']}")
            if pd.notna(row.get('Total Remaining (Amount)')):
                content_parts.append(f"Total Remaining: {row['Total Remaining (Amount)']}")
            if pd.notna(row.get('Event Name')):
                content_parts.append(f"Event: {row['Event Name']}")
            if pd.notna(row.get('Location')):
                content_parts.append(f"Location: {row['Location']}")
            if pd.notna(row.get('Attendance')):
                content_parts.append(f"Attendance: {row['Attendance']}")
            if pd.notna(row.get('Key Outcomes')):
                content_parts.append(f"Key Outcomes: {row['Key Outcomes']}")
            if pd.notna(row.get('Agenda')):
                content_parts.append(f"Agenda: {row['Agenda']}")
            if pd.notna(row.get('Decisions Made')):
                content_parts.append(f"Decisions Made: {row['Decisions Made']}")
            if pd.notna(row.get('Issue Title')):
                content_parts.append(f"Issue: {row['Issue Title']}")
            if pd.notna(row.get('Description')):
                content_parts.append(f"Description: {row['Description']}")
            if pd.notna(row.get('Comments')):
                content_parts.append(f"Comments: {row['Comments']}")
            
            return " | ".join(content_parts) if content_parts else f"Record ID: {row.get('ID', 'Unknown')}"
        
        df['document_content'] = df.apply(prepare_document_content, axis=1)
        
        # Create embeddings and vector store
        loader = DataFrameLoader(df, page_content_column='document_content')
        documents = loader.load()
        
        # Initialize embeddings with proper device handling and error catching
        try:
            import torch
            # Force CPU usage and avoid meta tensors
            torch.set_default_device('cpu')
            
            embeddings = HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-MiniLM-L6-v2",
                model_kwargs={'device': 'cpu'},
                encode_kwargs={'normalize_embeddings': True}
            )
        except Exception as embed_error:
            st.warning(f"Using alternative embedding method due to: {str(embed_error)}")
            # Fallback to a simpler embedding method
            from langchain_community.embeddings import FakeEmbeddings
            embeddings = FakeEmbeddings(size=384)
        
        vectorstore = FAISS.from_documents(documents, embeddings)
        # Configure retriever to get more context
        retriever = vectorstore.as_retriever(
            search_type="similarity",
            search_kwargs={"k": 10}  # Retrieve top 10 most relevant documents
        )
        
        # Create prompt templates
        report_template = """
        You are the Welfare Committee Secretary writing a warm, human, and inclusive report.
        
        CRITICAL INSTRUCTIONS:
        1. Use ONLY the data provided in the context below - DO NOT use general knowledge or make assumptions
        2. Write in first person as the secretary, using a warm and inclusive tone
        3. Write in flowing paragraphs and essay form - NO bullet points or lists
        4. Make it read like a story that connects all the welfare activities together
        5. Include ALL specific details from the data: names, amounts, dates, locations, attendance, outcomes
        
        STRUCTURE YOUR REPORT AS FLOWING PARAGRAPHS:
        
        Opening: Start with a warm introduction about the welfare committee's mission and period covered.
        
        Financial Overview: Write a detailed narrative paragraph about the finances. Mention total collections, how the money was spent, what remains, and what this means for the welfare of members. Make it conversational but include all the numbers.
        
        Events and Activities: Write in essay form about each event. Describe what happened, who attended, where it was held, and what outcomes were achieved. Connect the events to show how they contribute to member welfare. Make it feel inclusive and celebratory.
        
        Meetings and Decisions: Write flowing paragraphs about the meetings held. Discuss the agendas naturally, weave in the decisions made, and explain how these decisions impact the welfare of members. Make it read like you're telling a colleague about important discussions.
        
        Challenges and Progress: If there are issues or concerns in the data, discuss them in a supportive, solution-oriented way. Show empathy and commitment to addressing them.
        
        Closing: End with a warm, forward-looking statement about the committee's continued commitment to member welfare.
        
        TONE: Warm, professional, inclusive, human, and caring. Avoid corporate jargon. Write as if you're sharing good news with friends while maintaining professionalism.
        
        User request: {question}
        
        Welfare Committee Data:
        {context}
        
        Write your report in flowing essay form with connected paragraphs. NO bullet points. Make it human and inclusive.
        """
        
        normal_template = """
        You are the Welfare Committee Secretary answering a specific question in a warm, human way.
        
        CRITICAL INSTRUCTIONS:
        1. Use ONLY the data provided in the context below
        2. Write in a conversational, friendly tone while being professional
        3. Include specific details from the data: names, amounts, dates, locations
        4. Write in flowing sentences, not bullet points
        5. If the data doesn't contain the answer, say "I don't have that information in our current records"
        6. Make your response feel personal and caring, as befits a welfare committee
        
        User request: {question}
        
        Welfare Committee Data:
        {context}
        
        Answer warmly and conversationally using only the data above:
        """
        
        report_prompt = PromptTemplate(template=report_template, input_variables=["context", "question"])
        normal_prompt = PromptTemplate(template=normal_template, input_variables=["context", "question"])
        
        # Create QA chains with strict context usage
        report_chain = RetrievalQA.from_chain_type(
            llm=llm,
            retriever=retriever,
            chain_type="stuff",
            chain_type_kwargs={
                "prompt": report_prompt,
                "verbose": False
            },
            return_source_documents=False
        )
        
        normal_chain = RetrievalQA.from_chain_type(
            llm=llm,
            retriever=retriever,
            chain_type="stuff",
            chain_type_kwargs={
                "prompt": normal_prompt,
                "verbose": False
            },
            return_source_documents=False
        )
        
        return report_chain, normal_chain, df
    except Exception as e:
        st.error(f"Error loading data: {str(e)}")
        return None, None, None

def generate_response(query):
    report_chain, normal_chain, df = initialize_ai()
    
    if report_chain is None:
        return "Please ensure GOOGLE_API_KEY is set and SECRETARY FORM(1-6).xlsx file is available."
    
    if "report" in query.lower():
        chain_to_use = report_chain
    else:
        chain_to_use = normal_chain
    
    try:
        response = chain_to_use.run(query)
        return response
    except Exception as e:
        return f"An error occurred: {str(e)}"

def create_pdf_download(content):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, 
                          rightMargin=72, leftMargin=72,
                          topMargin=72, bottomMargin=72)
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'ReportTitle',
        parent=styles['Heading1'],
        fontSize=20,
        spaceAfter=30,
        alignment=TA_CENTER,
        textColor='#1a365d',
        fontName='Helvetica-Bold'
    )
    
    body_style = ParagraphStyle(
        'ReportBody',
        parent=styles['Normal'],
        fontSize=12,
        spaceAfter=14,
        alignment=TA_JUSTIFY,
        leftIndent=0,
        rightIndent=0,
        leading=16
    )
    
    story = []
    story.append(Paragraph("Welfare Committee Report", title_style))
    story.append(Spacer(1, 30))
    
    clean_content = content.replace("I am", "The committee").replace("I have", "The committee has")
    clean_content = clean_content.replace("my role", "the secretary's role")
    
    paragraphs = clean_content.split('\n\n')
    for para in paragraphs:
        if para.strip():
            clean_para = para.strip().replace('\n', ' ')
            story.append(Paragraph(clean_para, body_style))
            story.append(Spacer(1, 8))
    
    story.append(Spacer(1, 40))
    current_date = datetime.now().strftime("%B %d, %Y")
    story.append(Paragraph(f"Report Date: {current_date}", styles['Normal']))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def create_docx_download(content):
    buffer = io.BytesIO()
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    title = doc.add_heading('Welfare Committee Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    clean_content = content.replace("I am", "The committee").replace("I have", "The committee has")
    clean_content = clean_content.replace("my role", "the secretary's role")
    
    paragraphs = clean_content.split('\n\n')
    for para in paragraphs:
        if para.strip():
            clean_para = para.strip().replace('\n', ' ')
            p = doc.add_paragraph(clean_para)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    doc.add_paragraph()
    current_date = datetime.now().strftime("%B %d, %Y")
    footer_para = doc.add_paragraph(f"Report Date: {current_date}")
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# Header
st.markdown('<div class="main-header"><h1>🏛️ Welfare Secretary AI</h1><p>Your AI assistant for welfare committee management</p></div>', unsafe_allow_html=True)

# Initialize session state
if 'messages' not in st.session_state:
    st.session_state.messages = []

# Sidebar
st.sidebar.title("⚙️ Settings")
st.sidebar.markdown("---")

# File upload section
st.sidebar.subheader("📁 Data Management")
uploaded_file = st.sidebar.file_uploader("Upload Excel File", type=['xlsx', 'xls'])
if uploaded_file is not None:
    # Save uploaded file
    with open("SECRETARY FORM(1-6).xlsx", "wb") as f:
        f.write(uploaded_file.getbuffer())
    st.sidebar.success("✅ File uploaded!")
    # Clear cache to reload data
    st.cache_resource.clear()

st.sidebar.markdown("---")

# Status information
st.sidebar.subheader("📊 System Status")
report_chain, normal_chain, df = initialize_ai()
if df is not None:
    st.sidebar.success(f"✅ Data loaded: {len(df)} records")
    st.sidebar.info(f"📝 Columns: {len(df.columns)}")
else:
    st.sidebar.warning("⚠️ No data loaded")
    st.sidebar.info("Upload Excel file and set GOOGLE_API_KEY in secrets")

# Chat container
chat_container = st.container()

# Display chat messages
with chat_container:
    if len(st.session_state.messages) == 0:
        st.markdown("""
        <div style='text-align: center; padding: 3rem 1rem; color: #808080;'>
            <h2 style='color: #ffffff;'>👋 Welcome!</h2>
            <p style='font-size: 1.1rem; margin-top: 1rem;'>
                Start a conversation by typing a message below or use the quick actions.
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    for idx, message in enumerate(st.session_state.messages):
        if message["role"] == "user":
            st.markdown(f'''
            <div class="chat-message user-message">
                <div class="message-label user-label">You</div>
                <div class="message-content">{message["content"]}</div>
            </div>
            ''', unsafe_allow_html=True)
        else:
            st.markdown(f'''
            <div class="chat-message ai-message">
                <div class="message-label ai-label">🏛️ Welfare Secretary AI</div>
                <div class="message-content">{message["content"]}</div>
            </div>
            ''', unsafe_allow_html=True)
            
            # Add download buttons for AI responses
            col1, col2, col3 = st.columns([1, 1, 6])
            with col1:
                pdf_data = create_pdf_download(message["content"])
                st.download_button(
                    label="📄 PDF",
                    data=pdf_data,
                    file_name=f"welfare_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                    mime="application/pdf",
                    key=f"pdf_{idx}"
                )
            with col2:
                docx_data = create_docx_download(message["content"])
                st.download_button(
                    label="📝 Word",
                    data=docx_data,
                    file_name=f"welfare_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"docx_{idx}"
                )

# Quick Actions Section
st.markdown('<div class="quick-actions">', unsafe_allow_html=True)
st.markdown("### ⚡ Quick Actions")
col1, col2 = st.columns(2)
with col1:
    if st.button("📊 Generate Welfare Report", use_container_width=True):
        query = "Generate a comprehensive welfare report"
        st.session_state.messages.append({"role": "user", "content": query})
        with st.spinner("🤔 AI is thinking..."):
            response = generate_response(query)
            st.session_state.messages.append({"role": "assistant", "content": response})
        st.rerun()
    
    if st.button("💰 Financial Summary", use_container_width=True):
        query = "What are the total finances collected?"
        st.session_state.messages.append({"role": "user", "content": query})
        with st.spinner("🤔 AI is thinking..."):
            response = generate_response(query)
            st.session_state.messages.append({"role": "assistant", "content": response})
        st.rerun()

with col2:
    if st.button("📅 Recent Events", use_container_width=True):
        query = "Tell me about recent events"
        st.session_state.messages.append({"role": "user", "content": query})
        with st.spinner("🤔 AI is thinking..."):
            response = generate_response(query)
            st.session_state.messages.append({"role": "assistant", "content": response})
        st.rerun()
    
    if st.button("🤝 Meeting Summaries", use_container_width=True):
        query = "What meetings have been held?"
        st.session_state.messages.append({"role": "user", "content": query})
        with st.spinner("🤔 AI is thinking..."):
            response = generate_response(query)
            st.session_state.messages.append({"role": "assistant", "content": response})
        st.rerun()

st.markdown('</div>', unsafe_allow_html=True)

# Chat input at the bottom
st.markdown("---")
prompt = st.chat_input("💬 Ask about welfare committee activities...")

if prompt:
    # Add user message
    st.session_state.messages.append({"role": "user", "content": prompt})
    
    # Get AI response
    with st.spinner("🤔 AI is thinking..."):
        response = generate_response(prompt)
        st.session_state.messages.append({"role": "assistant", "content": response})
    
    st.rerun()